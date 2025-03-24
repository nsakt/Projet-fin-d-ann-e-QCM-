import os
import random
from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import *
from pathlib import Path



class Question :

    """Classe pour question"""

    def __init__ (self) :
        self.enonce = None
        self.reponses = []
        self.bonnerep = None
        self.bonnereptext = None

    def __len__(self):
        return len(self.reponses)
    


def create_enonce_eleve(chemin_liste, nb_questions):
    """Création d'une liste de questions a partir d'un fichier texte formaté comme décrit dans l'exercice, selon un nombre d'élèves"""

    with open(chemin_liste, 'r') as file:
        # Read each line in the file
        ListeQuestions = []
        
        compt_quest = 0
        compt_task = 1
        temp_enonce = None
        temp_reponses=[]
        temp_bonnerep = None
        temp_bonnereptext = None

        for line in file:
            if compt_quest < nb_questions :
                match compt_task:
                    case 1:
                        temp_enonce = line.strip()
                        compt_task += 1

                        print("DEBUG - Enoncé lu :", temp_enonce)


                    case 2 | 3 | 4 | 5:
                        temp_reponses.append(line.strip())
                        compt_task += 1
                        print("DEBUG - Réponses lue :", temp_reponses)

                    case 6:
                        temp_bonnerep = int(line.strip())
                        temp_bonnereptext = temp_reponses[temp_bonnerep-1]
                        compt_task += 1
                        print("DEBUG - Bonne réponse lue :", temp_bonnerep, "-", temp_bonnereptext)

                    case 7:
                        None
                        
                        TempQuestion = Question()
                        TempQuestion.enonce = temp_enonce
                        TempQuestion.reponses = temp_reponses
                        TempQuestion.bonnerep = temp_bonnerep
                        TempQuestion.bonnereptext = temp_bonnereptext

                        ListeQuestions.append(TempQuestion)

                        temp_enonce = None
                        temp_reponses=[]
                        temp_bonnerep = None
                        temp_bonnereptext = None

                        print("DEBUG - Question Enrgistrée - Enoncé :", ListeQuestions[compt_quest].enonce, " - Reponses: ", ListeQuestions[compt_quest].reponses, " - Bonne réponse : ", ListeQuestions[compt_quest].bonnerep, " ", ListeQuestions[compt_quest].bonnereptext)

                        compt_task = 1
                        compt_quest += 1

                        print("DEBUG - Liste total :", ListeQuestions)
                        
    return ListeQuestions

def shuffle_enonce_eleves(ListeBase):
    """Mélange l'énoncé d'un élève (une liste de questions) induviduels tout en gardant trace de la bonne réponse"""
    random.shuffle(ListeBase)
    for elt in ListeBase:
        random.shuffle(elt.reponses)
        for i in range(4):
            if elt.reponses[i] == elt.bonnereptext:
                elt.bonnerep = i+1
    return ListeBase


def create_enonce_eleves(chemin_liste, nb_questions, nb_eleve):
    """Crée une liste d'énoncés mélangés pour un nombre donné d'élèves, utilisant les fonctions précédentes"""
    Enonce_Final = []
    Liste_Questions_Temp = []
    for i in range(nb_eleve):
        Liste_Questions_Temp = create_enonce_eleve(chemin_liste,nb_questions)
        Liste_Questions_Temp = shuffle_enonce_eleves(Liste_Questions_Temp)
        Enonce_Final.append(Liste_Questions_Temp)
    return Enonce_Final


def create_correction_enonce(Enonce_Base):  
    """Crée une liste de set de réponses pour les énoncés donnés, sous forme de ints"""
    liste_rep_enonce=[]
    subliste_rep=[]
    for elt in Enonce_Base:
        for i in range(len(elt)):
            subliste_rep.append(elt[i].bonnerep)
        liste_rep_enonce.append(subliste_rep)
        subliste_rep=[]
    print(liste_rep_enonce)
    return liste_rep_enonce



def create_doc_correction(liste_corrections,nom_document):
    
    
    doc_corrige = Document()
    tempstring = ''
    cmpt_sujets=1
    cmpt_space=1
    for elt in liste_corrections :
        doc_corrige.add_paragraph('Sujet ' + str(cmpt_sujets))
        for i in range(len(elt)):
            match elt[i]:
                case 1:
                    tempstring += 'a'

                case 2:
                    tempstring += 'b'
                
                case 3:
                    tempstring += 'c'
                
                case 4:
                    tempstring += 'd'
            if cmpt_space == 5:
                tempstring += ' '
                cmpt_space = 1
            else:
                cmpt_space += 1
        
        cmpt_sujets += 1
        doc_corrige.add_paragraph(tempstring)
        tempstring = ''
        doc_corrige.add_paragraph(tempstring)
                
    doc_corrige.add_paragraph(tempstring)
    doc_corrige.save(nom_document + ".docx")






def create_doc_sujets_run(enonce_total,nom_document):
    

    tempstring = ''
    cmpt_sujets=1



    for elt in enonce_total:
        doc_sujet = Document()
        doc_sujet.add_paragraph("Réponses : ")
        table = doc_sujet.add_table(rows=2, cols=11)
        doc_sujet.add_paragraph("")
        table2 = doc_sujet.add_table(rows=2, cols=11)
        for i in range(10):
            cell = table.cell(0, i)
            if i < 5 :
                cell.text = str(i+1)
            else :
                cell = table.cell(0, i+1)
                cell.text = str(i+1)

        for i in range(10):
            cell = table2.cell(0, i)
            if i < 5 :
                cell.text = str(i+11)
            else :
                cell = table2.cell(0, i+1)
                cell.text = str(i+11)
        doc_sujet.add_paragraph("")

        run = doc_sujet.add_paragraph().add_run()
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(9)
        run.add_text('Nom et Prenom : ')
        run.add_break()
        run.add_text('Sujet  ' + str(cmpt_sujets))
        run.add_break()
        for i in range(len(elt)):
            run.add_text('Question : ' + str(i+1))
            run.add_break()
            run.add_text(elt[i].enonce)
            run.add_break()

            tempstring += ' | '
            
            for f in range(4):
                match f:
                    case 0 :
                        tempstring += 'a -'
                    
                    case 1 :
                        tempstring += 'b -'

                    case 2 :
                        tempstring += 'c -'

                    case 3 :
                        tempstring += 'd -'

                tempstring += str(elt[i].reponses[f]) + ' | '

            run.add_text(tempstring)
            run.add_break()
            run.add_break()
            tempstring = ''


        doc_sujet.save(nom_document + "_" + str(cmpt_sujets) +".docx")
        cmpt_sujets +=1
    
def programme_principal(chemin_fichier_questions, nb_eleves, nb_questions, nom_fichier_sortie, chemin_sortie):
    """Programme principal qui génère les sujets et les corrections"""
    Enonce_test = create_enonce_eleves(chemin_fichier_questions, nb_questions, nb_eleves)
    create_doc_sujets_run(Enonce_test, os.path.join(chemin_sortie, nom_fichier_sortie))
    Correction_Test = create_correction_enonce(Enonce_test)
    create_doc_correction(Correction_Test, os.path.join(chemin_sortie, nom_fichier_sortie + "_correction"))

def select_file():
    """Ouvre une boîte de dialogue pour sélectionner un fichier"""
    filename = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if filename:
        entry_file.delete(0, tk.END)
        entry_file.insert(0, filename)

def select_output_dir():
    """Ouvre une boîte de dialogue pour sélectionner un répertoire de sortie"""
    directory = filedialog.askdirectory()
    if directory:
        entry_output_dir.delete(0, tk.END)
        entry_output_dir.insert(0, directory)

def run_program():
    """Lance le programme principal avec les paramètres spécifiés"""
    chemin_fichier_questions = entry_file.get()
    nb_eleves = int(entry_nb_eleves.get())
    nb_questions = int(entry_nb_questions.get())
    nom_fichier_sortie = entry_output_name.get()
    chemin_sortie = entry_output_dir.get()

    if not chemin_fichier_questions or not nb_eleves or not nb_questions or not nom_fichier_sortie or not chemin_sortie:
        messagebox.showerror("Erreur", "Veuillez remplir tous les champs")
        return

    programme_principal(chemin_fichier_questions, nb_eleves, nb_questions, nom_fichier_sortie, chemin_sortie)
    messagebox.showinfo("Succès", "Les fichiers ont été générés avec succès")

# Création de l'interface graphique
root = tk.Tk()
root.title("Générateur de QCM")
current_directory = os.path.abspath(os.path.dirname(__file__))
icon_path = os.path.join(current_directory, "large_icon.ico")
if Path(icon_path).exists():
    root.iconbitmap(icon_path)

# Frame pour les entrées
frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

# Sélection du fichier de questions
tk.Label(frame, text="Fichier de questions :").grid(row=0, column=0, sticky="w")
entry_file = tk.Entry(frame, width=50)
entry_file.grid(row=0, column=1)
tk.Button(frame, text="Parcourir", command=select_file).grid(row=0, column=2)

# Nombre d'élèves
tk.Label(frame, text="Nombre d'élèves :").grid(row=1, column=0, sticky="w")
entry_nb_eleves = tk.Entry(frame, width=10)
entry_nb_eleves.grid(row=1, column=1, sticky="w")

# Nombre de questions
tk.Label(frame, text="Nombre de questions :").grid(row=2, column=0, sticky="w")
entry_nb_questions = tk.Entry(frame, width=10)
entry_nb_questions.grid(row=2, column=1, sticky="w")

# Nom du fichier de sortie
tk.Label(frame, text="Nom du fichier de sortie :").grid(row=3, column=0, sticky="w")
entry_output_name = tk.Entry(frame, width=50)
entry_output_name.grid(row=3, column=1)

# Répertoire de sortie
tk.Label(frame, text="Répertoire de sortie :").grid(row=4, column=0, sticky="w")
entry_output_dir = tk.Entry(frame, width=50)
entry_output_dir.grid(row=4, column=1)
tk.Button(frame, text="Parcourir", command=select_output_dir).grid(row=4, column=2)

# Bouton pour lancer le programme
tk.Button(frame, text="Générer les QCM", command=run_program).grid(row=5, column=1, pady=10)

# Lancement de l'interface
root.mainloop()

#zone de test

root.mainloop()

